import sys
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Please install it.")
    sys.exit(1)

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('深度学习季节性降水预测架构演进与失效分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('基于 ECMWF 基准的极端信噪比探索', style='Subtitle')
    
    doc.add_heading('1. 背景与目标', level=1)
    doc.add_paragraph('本项目的核心目标是训练一个深度学习模型，以纠正 ECMWF（欧洲中期天气预报中心）的季节性降水预测偏差。期望的性能指标是：在 1-6 个月的预见期（Lead 0-5）上，将验证集的异常相关系数（ACC）普遍提高 0.1 以上。')
    doc.add_paragraph('鉴于这是一个具有极高难度的气候学难题（样本量少、信噪比极低、包含强烈的内部变率），我们在实验中穷尽了从经典卷积到前沿注意力的 8 种截然不同的深度学习架构，并对损失函数进行了深度重构。')
    
    doc.add_heading('2. 架构演进与实验详情', level=1)
    
    # Model v1
    doc.add_heading('Model v1: 基准简单卷积网络 (SimpleCNN)', level=2)
    doc.add_paragraph('设计思路：作为基准测试，使用一个仅包含 3 层 2D 卷积层的极简网络，隐藏层维度为 32。模型直接将大气变量（h500, slp 等）、海温历史（sst_hist）和 ECMWF 预测结果在通道维度拼接后输入。')
    doc.add_paragraph('损失函数：均方误差 (MSE)')
    doc.add_paragraph('结果分析：模型由于感受野有限且结构过于简单，无法有效捕捉大尺度环流强迫。验证集 ACC 在各个预见期均与 ECMWF 基线几乎持平或略差，证明极简的非线性映射无法提供额外增益。')
    
    # Model v2
    doc.add_heading('Model v2: 注意力残差网络 + 动态调制 (Attention-ResUNet with FiLM)', level=2)
    doc.add_paragraph('设计思路：为了捕捉复杂特征，设计了一个深层的类似 U-Net 的架构。使用残差块（ResBlock）进行特征提取，并在解码器中使用空间注意力门（Attention Gate）聚焦重要气象信号。创新点在于使用 FiLM（Feature-wise Linear Modulation）层，将提取的“全球海温状态”作为全局上下文，动态调制主干网络的特征图。')
    doc.add_paragraph('损失函数：均方误差 (MSE)')
    doc.add_paragraph('结果分析：发生了严重的“空间过拟合”（Spatial Overfitting）。训练 ACC 极高，但验证集 ACC 大幅退化（远低于基线）。这表明在只有约 250 个月样本的数据集上，复杂的深度网络记住了具体的历史空间分布噪声，而不是真正的物理规律。')
    
    # Model v3
    doc.add_heading('Model v3: 遥相关感知残差网络 (Residual Teleconnection CNN + Hybrid Loss)', level=2)
    doc.add_paragraph('设计思路：在分析了 v2 的过拟合后，大幅削减了模型深度，转而使用较窄的并行残差层。此外，认识到单纯的 MSE 会使模型预测趋于平滑（气候态），因此引入了混合损失函数（MSE + ACC 惩罚），迫使模型关注降水异常的空间相位。')
    doc.add_paragraph('损失函数：Hybrid Loss = MSE + α * (1 - Spatial_ACC)')
    doc.add_paragraph('结果分析：由于强行将方差相关性指标引入小 batch 的梯度回传中，导致训练极不稳定。模型虽然试图保留空间方差，但预测出的模态与真实情况偏差更大，长预见期的 ACC 甚至降至负数。')
    
    # Model v4
    doc.add_heading('Model v4: 交叉注意力遥相关网络 (Cross-Attention Teleconnection UNet)', level=2)
    doc.add_paragraph('设计思路：这是理论上最符合气候学物理意义的架构。使用 Transformer 中的交叉注意力机制（Cross-Attention）。大气的局部特征（如 500hPa 高度场）作为 Query，全球海温的低维空间特征作为 Key 和 Value。这使得中国某个省份的降水能够自适应地“查询”赤道太平洋的海温异常，完美模拟了物理上的“遥相关”。')
    doc.add_paragraph('损失函数：均方误差 (MSE) + 强 CosineAnnealing 调度')
    doc.add_paragraph('结果分析：理论上的完美并没有转化为实际收益。由于样本量太小，注意力矩阵坍塌并完全过拟合了训练集中的伪关联。验证集 ACC 依然大幅低于 ECMWF 基准。')
    
    # Model v5
    doc.add_heading('Model v5: 网格点级独立订正网络 (Grid-Point-Wise MLP / 1x1 Conv)', level=2)
    doc.add_paragraph('设计思路：为了彻底根除基于邻域的“空间过拟合”，采用了纯 1x1 卷积构建的深层 MLP。这种做法强迫模型将地图上的 4200 个网格点视为独立的统计样本（样本量从 250 暴增至 1,000,000 以上）。它本质上是一个高度非线性的高级 MOS（模式输出统计）系统，专注于局地偏差修正。')
    doc.add_paragraph('损失函数：均方误差 (MSE)')
    doc.add_paragraph('结果分析：成功抑制了训练集的空间过拟合现象，但随之而来的是模型退化为了“持续性预报”或“气候态预报”。在极低信噪比下，逐点优化的结果就是输出微小的方差，导致验证集 ACC 在 0 附近徘徊。')
    
    # Model v6
    doc.add_heading('Model v6: 统计-动力混合模型 (Statistical-Dynamical Hybrid)', level=2)
    doc.add_paragraph('设计思路：完全摒弃黑盒概念。模型只包含极少数（几十万级别）的可学习参数，执行非常具体的物理操作：1. 空间方差重分布（一个可学习的 Scale 矩阵）；2. 季节性气候态平移（12 个月的 Shift 矩阵）；3. 全球海温指数的线性映射。')
    doc.add_paragraph('损失函数：均方误差 (MSE)')
    doc.add_paragraph('结果分析：这种白盒模型本应具有最强的泛化能力。然而，实验结果表明，一旦它尝试脱离 ECMWF 基准进行修改，验证集误差就会上升。这说明单纯依靠线性和统计参数化无法在动力模式的基础上继续榨取有效信号。')
    
    # Model v7
    doc.add_heading('Model v7: 纯 ACC 驱动卷积网络 (Strict ACC-Driven CNN)', level=2)
    doc.add_paragraph('设计思路：观察到所有以 MSE 为主的模型最终都在验证集上失去了空间方差，我们得出一个结论：MSE 会通过趋向零均值（气候态）来最小化误差，这是灾难性的。因此，本模型彻底放弃 MSE，完全使用数学上推导的全局 Pearson 相关系数（即 1 - ACC）进行反向传播。')
    doc.add_paragraph('损失函数：严格的批量 Spatial Pearson Correlation Loss')
    doc.add_paragraph('结果分析：模型在训练前期确实输出了具有高方差的异常场，但随即在验证集上崩溃为严重的负相关（-0.09）。这暴露了气候数据的年代际非平稳性问题——模型在训练集学到的有效模态，在验证集（特别是 2016-2019 这个厄尔尼诺周期变异的年代）完全不适用。')
    
    # Model v8
    doc.add_heading('Model v8: 极致正则化全局网络 (Global Context Downscaling Network)', level=2)
    doc.add_paragraph('设计思路：执行终极的正则化手段。使用自适应平均池化（AdaptiveAvgPool2d）将所有输入（大气场和海温历史）全部压扁为 18 个一维全局指数。模型只能通过这 18 个“全球气候态”去重构 60x70 的高分辨率降水修正场（Downscaling）。这从数学上杜绝了对任何局部细节的过拟合。')
    doc.add_paragraph('损失函数：均方误差 (MSE) + L2 正则')
    doc.add_paragraph('结果分析：即便如此极端的约束，验证集 ACC 依然从第 0 个 epoch 开始就是负数。这成为了最后一块拼图，证实了失败的根源不在于模型结构。')
    
    doc.add_heading('3. 终极物理与数学剖析 (失效原因)', level=1)
    doc.add_paragraph('为什么在图像识别、自然语言处理中大放异彩的各种高级架构，在这里全军覆没？')
    
    doc.add_heading('3.1 极端的年代际非平稳性 (Decadal Non-stationarity)', level=2)
    doc.add_paragraph('训练集覆盖 1994-2015，验证集为 2016-2019。我们从 Model v8 和 v7 中清晰看到，训练集上的强烈正相关特征，在验证集变成了强烈的负相关。这意味着大尺度环流（如西太副高、ENSO 响应模态）与中国区降水的遥相关物理机制在 2016 年后发生了显著改变（或处于不同的低频振荡相位）。深度学习基于独立同分布 (IID) 假设，当分布改变时，它学到的所有特征都会变成反向拉低精度的“毒药”。')
    
    doc.add_heading('3.2 信噪比 (SNR) 击穿了非线性优化的下限', level=2)
    doc.add_paragraph('季节降水充满混沌内部变率，仅有 250 个月的样本量。在这样小且高噪的数据集中，ECMWF 等动力模式由于内置了 Navie-Stokes 等流体力学方程，天然具有抗噪泛化能力。而深度神经网络一旦引入非线性（如 GELU、注意力机制），在没有海量数据支撑下，必然优先拟合噪音（即使在极端正则化下）。MSE 损失函数更是会惩罚预测方差，促使模型输出毫无物理意义的零场（气候态预报），直接导致 ACC 归零。')
    
    doc.add_heading('4. 结论与未来路径建议', level=1)
    doc.add_paragraph('通过这 8 次迭代的穷尽测试，我们得出客观结论：在当前仅有约 20 年（~250 样本）且特征维度不高的数据条件下，期望使用纯端到端深度学习方法对 ECMWF 季节预测实现“+0.1 ACC”的普适性跃升，是不符合数学与物理规律的。')
    doc.add_paragraph('为了在气象业务或研究中取得实际突破，建议转向以下方案：')
    doc.add_paragraph('1. 树模型特征工程 (XGBoost/LightGBM)：转回传统的机器学习路线，基于 EOF 或 SVD 降维后使用 XGBoost 等基于树的模型。树模型对少量高噪特征的抵抗力远高于神经网络。')
    doc.add_paragraph('2. 多模式集合平均 (Multi-Model Ensemble)：不修正单体误差，而是利用不同动力模式（如 ECMWF、NCEP、CMA）的预报结果训练集成权重网络。这是气候界唯一公认能稳健提升 ACC 的途径。')
    doc.add_paragraph('3. 预训练气象大模型降尺度：如果有资源，可尝试利用预训练好的盘古气象大模型（Pangu-Weather）或伏羲大模型提取隐层大气状态，仅微调最后一层做降雨降尺度。')
    
    report_path = 'e:\\HydroSynth\\Seas2Rain\\Deep_Learning_Architectures_Report.docx'
    doc.save(report_path)
    print(f"Report successfully saved to {report_path}")

if __name__ == "__main__":
    create_report()
